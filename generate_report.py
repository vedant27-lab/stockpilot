from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "WTL_Mini_Project_Report_StockPilot_Inventory.docx"
SCREENSHOTS = [
    ROOT / "image.png",
    ROOT / "image copy.png",
    ROOT / "image copy 2.png",
]


def set_page_borders(section):
    sect_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")

    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "24")
        element.set(qn("w:color"), "auto")
        pg_borders.append(element)

    sect_pr.append(pg_borders)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size in ((1, 14), (2, 12)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    if "Body Label" not in doc.styles:
        style = doc.styles.add_style("Body Label", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)

    if "Code Note" not in doc.styles:
        style = doc.styles.add_style("Code Note", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)

    if "Code Block" not in doc.styles:
        style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(9.5)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.0


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.add_run(text).bold = True


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc, title, code):
    doc.add_paragraph(title, style="Body Label")
    paragraph = doc.add_paragraph(style="Code Block")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.15)

    lines = code.strip("\n").splitlines()
    for index, line in enumerate(lines):
        run = paragraph.add_run(line.rstrip())
        if index < len(lines) - 1:
            run.add_break()


def extract_snippet(text, start_marker, end_marker):
    start = text.index(start_marker)
    end = text.index(end_marker, start)
    return text[start:end].rstrip()


def add_file_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "File Name"
    header[1].text = "Extension"
    header[2].text = "Purpose"

    for cell in header:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if paragraph.runs:
                paragraph.runs[0].bold = True
        set_cell_shading(cell, "D9EAF7")

    rows = [
        ("index.html", ".html", "Main inventory workspace layout"),
        ("style.css", ".css", "Smooth responsive styling, theme system, chart and card styles"),
        ("script.js", ".js", "Frontend logic for inventory, analytics, charts, and access sharing"),
        ("server.js", ".js", "Local Node.js backend server and API routing"),
        ("api/index.js", ".js", "Vercel-compatible serverless API entry"),
        ("lib/store.js", ".js", "Seed data generation and shared data store utilities"),
        ("data/store.json", ".json", "Persistent demo dataset for inventory, movements, and shares"),
        ("package.json", ".json", "Project metadata and start script"),
        ("vercel.json", ".json", "Vercel deployment rewrite configuration"),
        ("README.md", ".md", "Project setup and deployment notes"),
    ]

    for name, ext, purpose in rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = ext
        row[2].text = purpose
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_document():
    doc = Document()
    configure_styles(doc)

    script_text = (ROOT / "script.js").read_text(encoding="utf-8")
    server_text = (ROOT / "server.js").read_text(encoding="utf-8")
    index_text = (ROOT / "index.html").read_text(encoding="utf-8")
    store_text = (ROOT / "lib" / "store.js").read_text(encoding="utf-8")

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        set_page_borders(section)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("WTL Mini-project")
    title_run.bold = True
    title_run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    subrun = subtitle.add_run("Project Report: StockPilot Inventory Management System")
    subrun.italic = True
    subrun.font.size = Pt(12)

    overview = doc.add_table(rows=4, cols=2)
    overview.style = "Table Grid"
    items = [
        ("Project Title", "StockPilot Inventory Management System"),
        ("Project Type", "Dynamic web application for inventory operations"),
        ("Repository", "https://github.com/vedant27-lab/stockpilot.git"),
        ("Prepared For", "WTL Mini-project submission"),
    ]
    for row, (label, value) in zip(overview.rows, items):
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row.cells[0], "EAF2F8")

    doc.add_paragraph("")

    add_heading(doc, "Problem Statement of Mini-Project")
    doc.add_paragraph(
        "Retail stores and small businesses often struggle to maintain accurate stock records, "
        "track incoming and outgoing inventory, identify reorder requirements, and coordinate access "
        "among team members. Manual methods lead to stock mismatch, delayed updates, and poor visibility. "
        "This project solves that problem by providing a dynamic inventory management web application with "
        "analytics, stock movement tracking, and controlled access sharing."
    )

    add_heading(doc, "Objective")
    doc.add_paragraph(
        "To learn web technologies like HTML, CSS, JS, JSP/SERVELET/ASP.NET/PHP, MYSQL, XML, etc. "
        "to implement dynamic web application."
    )
    doc.add_paragraph(
        "The project applies these concepts by building a responsive inventory workspace using HTML, CSS, "
        "JavaScript, Node.js, and JSON-based persistence."
    )

    add_heading(doc, "Technologies Used")
    doc.add_paragraph("S/W:", style="Body Label")
    add_bullets(
        doc,
        [
            "Visual Studio Code",
            "HTML5",
            "CSS3",
            "Vanilla JavaScript",
            "Node.js built-in HTTP server",
            "JSON file storage",
            "Git and GitHub",
            "Google Chrome / Microsoft Edge",
            "Vercel-compatible API structure",
        ],
    )
    doc.add_paragraph("H/W:", style="Body Label")
    add_bullets(
        doc,
        [
            "Laptop or desktop system",
            "Intel/AMD processor",
            "4 GB RAM or above",
            "Minimum 200 MB free storage",
            "Internet connection for repository and hosting access",
        ],
    )

    add_heading(doc, "Introduction of Topic/Project")
    doc.add_paragraph(
        "StockPilot is a modern inventory management web application designed for small business use. "
        "The system maintains stock items, records stock-in and stock-out activity, shows low-stock alerts, "
        "presents category-wise chart analysis, and supports shared workspace access for multiple roles. "
        "Its scope includes inventory visibility, operational analytics, responsive design, smooth user experience, "
        "and simple collaboration. The project is important because it addresses practical stock-management needs "
        "through a clear, minimal, and easy-to-use digital system."
    )

    add_heading(doc, "Project Architecture")
    doc.add_paragraph(
        "The project follows a lightweight client-server architecture. The frontend layer renders the interface "
        "and interacts with backend APIs. The backend validates requests, updates inventory records, manages demo "
        "seed data, and serves the application files. The application also supports a Vercel-compatible API entry "
        "for cloud deployment."
    )
    add_bullets(
        doc,
        [
            "Presentation Layer: index.html and style.css build the inventory dashboard and responsive UI.",
            "Client Logic Layer: script.js handles rendering, analytics calculations, chart drawing, theme switching, and API communication.",
            "Server Layer: server.js serves local API routes and static assets.",
            "Deployment API Layer: api/index.js provides a serverless-compatible backend entry for Vercel.",
            "Data Layer: lib/store.js and data/store.json maintain inventory items, stock movements, and shared-access data.",
            "Version Control Layer: the project is stored in the GitHub repository https://github.com/vedant27-lab/stockpilot.git",
        ],
    )

    add_heading(doc, "Project design/functionalities/features")
    doc.add_paragraph("Module 1: Inventory Item Management", style="Body Label")
    add_bullets(
        doc,
        [
            "Add new inventory items with name, SKU, category, supplier, price, opening quantity, and reorder level",
            "Display all items in a searchable inventory register",
            "Delete inventory items when required",
            "Show low-stock status using clear visual indicators",
        ],
    )
    doc.add_paragraph("Module 2: Stock Movement Management", style="Body Label")
    add_bullets(
        doc,
        [
            "Record stock-in and stock-out operations",
            "Automatically update current quantity after each movement",
            "Store note and time for each movement event",
            "Maintain recent inventory activity timeline",
        ],
    )
    doc.add_paragraph("Module 3: Analytics and Graph Analysis", style="Body Label")
    add_bullets(
        doc,
        [
            "Show summary cards for total items, total units, inventory value, and reorder count",
            "Visualize category-wise stock distribution using an SVG bar chart",
            "Show aggregate movement insights such as incoming and outgoing units",
            "Use richer demo data to create meaningful graphs and screenshots",
        ],
    )
    doc.add_paragraph("Module 4: Access Sharing Module", style="Body Label")
    add_bullets(
        doc,
        [
            "Create access invites for teammates",
            "Assign Viewer or Editor role",
            "Show invite token and status in the shared-access panel",
            "Revoke access records when needed",
        ],
    )
    doc.add_paragraph("Module 5: UI and Accessibility Module", style="Body Label")
    add_bullets(
        doc,
        [
            "Smooth and responsive layout for desktop and mobile devices",
            "Light and dark mode toggle at the corner",
            "Accessible labels, focus states, skip link, and semantic sections",
            "Mobile-friendly table-to-card transformation on smaller screens",
        ],
    )

    add_heading(doc, "Working/flowchart/algorithms")
    doc.add_paragraph("Overall Working:", style="Body Label")
    add_numbered(
        doc,
        [
            "User opens the StockPilot inventory workspace.",
            "Frontend requests dashboard data from the backend API.",
            "Backend loads products, stock movements, and shared-access records from the store.",
            "Frontend renders summary cards, graphs, inventory table, alerts, timeline, and share list.",
            "When a new inventory item is added, the frontend sends the form data to the products API.",
            "When stock movement is recorded, the backend validates the quantity and updates stock accordingly.",
            "When a share invite is created, the backend stores the access record with role and token.",
            "Frontend reloads dashboard data and updates charts and tables immediately.",
        ],
    )
    doc.add_paragraph("Textual Flowchart:", style="Body Label")
    add_bullets(
        doc,
        [
            "Start",
            "Open inventory workspace",
            "Fetch dashboard data",
            "Render items, analytics, and access records",
            "User selects action: add item / record movement / create invite / search / delete / reset / theme switch",
            "Validate request",
            "Update store",
            "Refresh dashboard",
            "Stop",
        ],
    )
    doc.add_paragraph("Algorithm for Stock Movement:", style="Body Label")
    add_numbered(
        doc,
        [
            "Accept selected item, movement type, quantity, and note from the user.",
            "Check whether the selected item exists.",
            "Validate that quantity is greater than zero.",
            "If movement type is stock-out, verify stock availability.",
            "Update quantity by adding or subtracting the movement amount.",
            "Create a movement record with note, time, and amount value.",
            "Save updated data and refresh the dashboard.",
        ],
    )

    add_heading(doc, "Other information")
    add_bullets(
        doc,
        [
            "The project includes a large 70-item demo inventory dataset for realistic testing and screenshots.",
            "The chart is generated without external chart libraries, using SVG and JavaScript.",
            "The app supports both local Node.js execution and a Vercel-compatible API route structure.",
            "The GitHub repository used for the project is https://github.com/vedant27-lab/stockpilot.git",
        ],
    )

    add_heading(doc, "Conclusion")
    doc.add_paragraph(
        "StockPilot Inventory Management System successfully demonstrates the implementation of a dynamic web application "
        "for business use. The application combines inventory control, stock movement tracking, graph-based analytics, "
        "responsive UI, and access-sharing features in one practical system. It clearly reflects the learning and application "
        "of core web technologies in a real-world mini-project format."
    )

    doc.add_page_break()

    add_heading(doc, "Code snippets explaining features and architecture")
    doc.add_paragraph(
        "The following code snippets are taken from the current project and show how the major features are implemented.",
        style="Code Note",
    )

    add_code_block(
        doc,
        "Snippet 1: Theme toggle handling from script.js",
        extract_snippet(
            script_text,
            "function resolvePreferredTheme() {",
            'els.productForm.addEventListener("submit", addProduct);',
        ),
    )
    doc.add_paragraph(
        "Explanation: This code stores and restores the selected theme, applies the theme to the document root, "
        "and keeps the floating corner toggle synchronized with the current mode.",
        style="Code Note",
    )

    add_code_block(
        doc,
        "Snippet 2: Inventory analytics calculation from script.js",
        extract_snippet(
            script_text,
            "function getAnalytics() {",
            "function renderStats() {",
        ),
    )
    doc.add_paragraph(
        "Explanation: This snippet calculates total units, inventory value, reorder alerts, incoming stock, outgoing stock, "
        "and category-wise data required for graph analysis.",
        style="Code Note",
    )

    add_code_block(
        doc,
        "Snippet 3: Category chart rendering from script.js",
        extract_snippet(
            script_text,
            "function renderCategoryChart() {",
            "function renderAll() {",
        ),
    )
    doc.add_paragraph(
        "Explanation: This feature generates the category stock bar chart using SVG elements, without any complex framework or chart library.",
        style="Code Note",
    )

    add_code_block(
        doc,
        "Snippet 4: Inventory movement API handling from server.js",
        extract_snippet(
            server_text,
            '  if (req.method === "POST" && url.pathname === "/api/movements") {',
            '  if (req.method === "POST" && url.pathname === "/api/shares") {',
        ),
    )
    doc.add_paragraph(
        "Explanation: This backend route validates stock movement requests, ensures stock cannot go negative, updates the item quantity, "
        "and stores a movement record for the timeline and analytics.",
        style="Code Note",
    )

    add_code_block(
        doc,
        "Snippet 5: Seed catalog generation from lib/store.js",
        extract_snippet(
            store_text,
            "function createSeedState() {",
            "let memoryStore = createSeedState();",
        ),
    )
    doc.add_paragraph(
        "Explanation: This section creates the rich demo dataset used to populate the application with many inventory items, "
        "movement events, and share records for testing and screenshots.",
        style="Code Note",
    )

    add_code_block(
        doc,
        "Snippet 6: Main interface structure from index.html",
        extract_snippet(
            index_text,
            "<body>",
            '  <script src="script.js"></script>',
        ),
    )
    doc.add_paragraph(
        "Explanation: This markup defines the major UI architecture including the hero section, summary cards, analytics panel, "
        "inventory forms, register table, movement timeline, and access-sharing panel.",
        style="Code Note",
    )

    add_heading(doc, "Code")
    doc.add_paragraph("All code files with their names and proper extensions are listed below.", style="Code Note")
    add_file_table(doc)

    add_heading(doc, "Output/screenshots")
    doc.add_paragraph(
        "The screenshots below are taken from the current inventory-management version of the website.",
        style="Code Note",
    )

    figure = 1
    for image_path in SCREENSHOTS:
        if image_path.exists():
            doc.add_paragraph(f"Screenshot: {image_path.name}", style="Body Label")
            doc.add_picture(str(image_path), width=Inches(5.9))
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.add_run(f"Figure {figure}: Website screenshot").italic = True
            figure += 1

    doc.add_paragraph("Suggested screenshot coverage:", style="Body Label")
    add_numbered(
        doc,
        [
            "Dashboard with summary cards and chart analysis",
            "Inventory add form and stock movement form",
            "Inventory register with many items",
            "Reorder alert section",
            "Recent movement timeline",
            "Access-sharing panel with role entries",
            "Dark mode screen",
            "Mobile responsive screen",
        ],
    )

    note = doc.add_paragraph()
    note_run = note.add_run("***note: for this assignment, no write-up, just prepare this report and print it.***")
    note_run.bold = True

    final_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_borders(final_section)
    add_heading(doc, "Repository Reference")
    doc.add_paragraph("GitHub Repository: https://github.com/vedant27-lab/stockpilot.git")
    doc.add_paragraph("Local Project Folder: WTL Mini-project")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
